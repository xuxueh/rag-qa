"""生成测试文档（txt/md/pdf/docx 各一份），用于验证 smart_loader"""
import os

def make_test_docs(doc_dir: str):
    os.makedirs(doc_dir, exist_ok=True)

    # TXT
    with open(os.path.join(doc_dir, "01-测试制度.txt"), "w", encoding="utf-8") as f:
        f.write("公司测试制度\n第一条：测试制度用于验证多格式加载。\n第二条：TXT 文件是最基础的格式。\n")

    # Markdown
    with open(os.path.join(doc_dir, "02-测试说明.md"), "w", encoding="utf-8") as f:
        f.write("# 测试说明\n\nMarkdown 文件支持标题和列表：\n\n- 第一条：Markdown 直接按文本读取\n- 第二条：适合技术文档\n")

    # PDF（用 reportlab 生成中文 PDF）
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(os.path.join(doc_dir, "03-测试手册.pdf"))
    c.setFont("Helvetica", 12)
    c.drawString(72, 800, "Test Document - RAG QA System")
    c.drawString(72, 780, "This PDF is for format testing.")
    c.showPage()
    c.save()

    # Word
    from docx import Document
    doc = Document()
    doc.add_heading("测试 Word 文档", 0)
    doc.add_paragraph("第一条：Word 文档通过 python-docx 读取文本内容。")
    doc.add_paragraph("第二条：用于验证 docx 格式的解析。")
    doc.save(os.path.join(doc_dir, "04-测试文档.docx"))

    print(f"✅ 测试文档已生成到 {doc_dir}")
