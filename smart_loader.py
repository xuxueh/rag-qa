"""
smart_loader.py - 多格式文档加载器（TXT / Markdown / PDF / Word）
=================================================================
按扩展名自动选择解析方式：
- .txt / .md  → TextLoader（纯文本/ Markdown 直接读）
- .pdf       → PyPDFLoader（pypdf 解析，带页码 metadata）
- .docx      → Docx2txtLoader（Word 文本提取）

用法：from smart_loader import load_documents
"""
import os
import sys

sys.stdout.reconfigure(encoding="utf-8")


def load_documents(doc_dir: str):
    """按扩展名加载目录下所有支持的文档，返回 Document 列表"""
    from langchain_community.document_loaders import TextLoader

    docs = []
    files = sorted(os.listdir(doc_dir))
    for fname in files:
        path = os.path.join(doc_dir, fname)
        if not os.path.isfile(path):
            continue
        ext = fname.lower().rsplit(".", 1)[-1] if "." in fname else ""

        try:
            if ext in ("txt", "md"):
                docs.extend(TextLoader(path, encoding="utf-8").load())
            elif ext == "pdf":
                from langchain_community.document_loaders import PyPDFLoader
                docs.extend(PyPDFLoader(path).load())
            elif ext == "docx":
                # 用 python-docx 解析（不依赖 docx2txt）
                from docx import Document as DocxDocument
                from langchain_core.documents import Document
                d = DocxDocument(path)
                text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
                docs.append(Document(page_content=text, metadata={"source": path}))
            else:
                print(f"⚠️ 跳过不支持的文件: {fname}")
        except Exception as e:
            print(f"⚠️ 加载失败 {fname}: {e}")

    print(f"✓ 加载 {len(docs)} 份文档（{len(files)} 个文件）")
    return docs


if __name__ == "__main__":
    # 自测：加载测试目录（txt/md/pdf/docx 混合）
    test_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "test_docs")
    if not os.path.exists(test_dir):
        print("测试目录不存在，先生成测试文件...")
        from make_test_docs import make_test_docs
        make_test_docs(test_dir)
    docs = load_documents(test_dir)
    for d in docs[:10]:
        src = d.metadata.get("source", "?")
        print(f"  [{src.split(chr(92))[-1].split('/')[-1]}] {d.page_content[:30]}...")
